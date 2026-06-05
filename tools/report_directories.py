import argparse
import json
import re
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, Twips


if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    sys.stderr.reconfigure(encoding="utf-8", errors="replace")


TITLE_TOC = "目录"
TITLE_FIGURES = "插图清单"
TITLE_TABLES = "附表清单"
TITLE_INTRODUCTION = "引言"
INTRO_STYLE_NAME = "Report Intro Heading"
INTRO_SUBHEADING_STYLE_NAME = "Report Intro Subheading"

HEADING_TEXT_RE = re.compile(r"^\d+(?:\.\d+){0,2}\s+\S+")
NUMBERED_HEADING_RE = re.compile(r"^(?P<number>\d+(?:\.\d+){0,2})\s+(?P<title>\S.*)$")
FIG_RE = re.compile(r"^图\s+\d+(?:\.\d+)?\s+")
TAB_RE = re.compile(r"^表\s+\d+(?:\.\d+)?\s+")
TOC_LINE_RE = re.compile(r"^(\d+(?:\.\d+){0,2})\s+(.+?)\s+(\d+)$")
HEADING_STYLE_IDS = {1: "Heading1", 2: "Heading2", 3: "Heading3"}
HEADING_STYLE_NAMES = {
    1: ("Heading 1", "标题 1"),
    2: ("Heading 2", "标题 2"),
    3: ("Heading 3", "标题 3"),
}
DEFAULT_FORMAT_CONFIG = {
    "titles": {
        "toc": "目录",
        "figures": "插图清单",
        "tables": "附表清单",
        "introduction": "引言",
    },
    "page_breaks": {
        "before_figures": True,
        "before_tables": True,
        "before_body": True,
    },
    "front_title": {"font": "仿宋", "size": 16, "bold": True},
    "toc_1": {"font": "宋体", "size": 12, "bold": None},
    "toc_2": {"font": "宋体", "size": 12, "bold": None},
    "toc_3": {"font": "宋体", "size": 12, "bold": None},
    "heading_1": {"font": "宋体", "size": 14, "bold": None},
    "heading_2": {"font": "宋体", "size": 12, "bold": None},
    "heading_3": {"font": "宋体", "size": 12, "bold": None},
    "intro": {"font": "宋体", "size": 14, "bold": True},
    "intro_subheading": {"font": "宋体", "size": 12, "bold": True},
    "caption": {"font": "宋体", "size": 12, "bold": None},
    "list_entry": {"font": "宋体", "size": 12, "bold": None},
    "line_spacing": 1.5,
}
FORMAT_CONFIG = DEFAULT_FORMAT_CONFIG


def merge_format_config(base: dict, override: dict | None) -> dict:
    merged = {}
    for key, value in base.items():
        if isinstance(value, dict):
            item = dict(value)
            if isinstance(override, dict) and isinstance(override.get(key), dict):
                item.update(override[key])
            merged[key] = item
        else:
            merged[key] = override.get(key, value) if isinstance(override, dict) else value
    if isinstance(override, dict):
        for key, value in override.items():
            if key not in merged:
                merged[key] = value
    return merged


def load_format_config(path: Path | None) -> dict:
    if path is None:
        return DEFAULT_FORMAT_CONFIG
    if not path.exists():
        print(f"format_config_missing={path}")
        return DEFAULT_FORMAT_CONFIG
    with path.open("r", encoding="utf-8-sig") as handle:
        data = json.load(handle)
    if not isinstance(data, dict):
        raise ValueError("Format config must be a JSON object.")
    return merge_format_config(DEFAULT_FORMAT_CONFIG, data)


def set_format_config(config: dict) -> None:
    global FORMAT_CONFIG
    FORMAT_CONFIG = config


def format_option(name: str, key: str):
    section = FORMAT_CONFIG.get(name, {})
    if isinstance(section, dict) and key in section:
        return section[key]
    default_section = DEFAULT_FORMAT_CONFIG.get(name, {})
    if isinstance(default_section, dict):
        return default_section.get(key)
    return None


def line_spacing() -> float:
    value = FORMAT_CONFIG.get("line_spacing", DEFAULT_FORMAT_CONFIG["line_spacing"])
    return float(value)


def config_title(name: str) -> str:
    titles = FORMAT_CONFIG.get("titles", {})
    if isinstance(titles, dict) and name in titles:
        return str(titles[name])
    return str(DEFAULT_FORMAT_CONFIG["titles"][name])


def page_break_enabled(name: str) -> bool:
    page_breaks = FORMAT_CONFIG.get("page_breaks", {})
    if isinstance(page_breaks, dict) and name in page_breaks:
        return bool(page_breaks[name])
    return bool(DEFAULT_FORMAT_CONFIG["page_breaks"][name])


def apply_run_format(run, style_key: str) -> None:
    ensure_run_fonts(
        run,
        str(format_option(style_key, "font")),
        format_option(style_key, "size"),
        format_option(style_key, "bold"),
    )


def apply_style_format(style, style_key: str) -> None:
    ensure_style_fonts(
        style,
        str(format_option(style_key, "font")),
        format_option(style_key, "size"),
        format_option(style_key, "bold"),
    )


def clean(text: str) -> str:
    return " ".join(text.split())


def title_toc() -> str:
    return config_title("toc")


def title_figures() -> str:
    return config_title("figures")


def title_tables() -> str:
    return config_title("tables")


def title_introduction() -> str:
    return config_title("introduction")


def front_titles() -> set[str]:
    return {title_toc(), title_figures(), title_tables()}


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def ensure_run_fonts(run, font_name: str, size_pt: float | None = None, bold: bool | None = None):
    run.font.name = font_name
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(key), font_name)


def ensure_style_fonts(style, font_name: str, size_pt: float | None = None, bold: bool | None = None):
    style.font.name = font_name
    if size_pt is not None:
        style.font.size = Pt(size_pt)
    if bold is not None:
        style.font.bold = bold
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(key), font_name)


def set_paragraph_style_id(paragraph, style_id: str) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    pstyle = ppr.pStyle
    if pstyle is None:
        pstyle = OxmlElement("w:pStyle")
        ppr.insert(0, pstyle)
    pstyle.set(qn("w:val"), style_id)


def clear_outline_level(paragraph) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    for outline in list(ppr.xpath("./w:outlineLvl")):
        ppr.remove(outline)


def set_outline_level(paragraph, level: int) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    for outline in list(ppr.xpath("./w:outlineLvl")):
        ppr.remove(outline)
    outline = OxmlElement("w:outlineLvl")
    outline.set(qn("w:val"), str(level - 1))
    ppr.append(outline)


def suppress_numbering(paragraph) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    for num_pr in list(ppr.xpath("./w:numPr")):
        ppr.remove(num_pr)
    num_pr = OxmlElement("w:numPr")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "0")
    num_pr.append(num_id)
    ppr.append(num_pr)


def ensure_page_break_at_start(paragraph) -> None:
    for br in paragraph._p.xpath('.//w:br[@w:type="page"]'):
        return
    first_run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    first_run._r.insert(0, br)


def remove_page_breaks(paragraph) -> None:
    for br in list(paragraph._p.xpath('.//w:br[@w:type="page"]')):
        br.getparent().remove(br)


def replace_paragraph_text(paragraph, text: str) -> None:
    had_page_break = bool(paragraph._p.xpath('.//w:br[@w:type="page"]'))
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    paragraph.add_run(text)
    if had_page_break:
        ensure_page_break_at_start(paragraph)


def find_style(doc: Document, names=(), style_ids=()):
    name_set = {name.casefold() for name in names}
    id_set = set(style_ids)
    for style in doc.styles:
        if style.name and style.name.casefold() in name_set:
            return style
        if getattr(style, "style_id", None) in id_set:
            return style
    return None


def ensure_style(doc: Document, name: str, base: str = "Normal"):
    style = find_style(doc, names=(name,))
    if style is None:
        style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    base_style = find_style(doc, names=(base, "Normal", "正文"), style_ids=("Normal",))
    if base_style is not None:
        style.base_style = base_style
    apply_style_format(style, "caption")
    style.paragraph_format.line_spacing = line_spacing()
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    return style


def ensure_table_of_figures_style(doc: Document):
    style = find_style(doc, names=("table of figures",))
    if style is None:
        style = doc.styles.add_style("table of figures", WD_STYLE_TYPE.PARAGRAPH)
    base_style = find_style(doc, names=("Normal", "正文"), style_ids=("a", "Normal"))
    if base_style is not None:
        style.base_style = base_style
    apply_style_format(style, "list_entry")
    style.paragraph_format.line_spacing = line_spacing()
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    return style


def ensure_toc_styles(doc: Document) -> None:
    for level in (1, 2, 3):
        style = find_style(doc, names=(f"toc {level}", f"TOC {level}"))
        if style is None:
            continue
        apply_style_format(style, f"toc_{level}")
        style.paragraph_format.line_spacing = line_spacing()
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)


def find_paragraph(doc: Document, text: str, start: int = 0):
    for i, paragraph in enumerate(doc.paragraphs[start:], start):
        if clean(paragraph.text) == text:
            return i
    return None


def remove_paragraphs_before(doc: Document, text: str) -> int:
    idx = find_paragraph(doc, text)
    if idx is None or idx == 0:
        return 0
    for paragraph in list(doc.paragraphs[:idx]):
        remove_paragraph(paragraph)
    return idx


def style_name(paragraph) -> str:
    style = paragraph.style
    if style is None:
        return ""
    return style.name or ""


def style_id(paragraph) -> str:
    style = paragraph.style
    if style is None:
        return ""
    return getattr(style, "style_id", "") or ""


def is_heading_style(paragraph, levels=(1, 2, 3)) -> bool:
    name = style_name(paragraph)
    sid = style_id(paragraph)
    for level in levels:
        if sid == HEADING_STYLE_IDS[level] or name in HEADING_STYLE_NAMES[level]:
            return True
    return False


def get_heading_style(doc: Document, level: int):
    style = find_style(
        doc,
        names=HEADING_STYLE_NAMES[level],
        style_ids=(HEADING_STYLE_IDS[level],),
    )
    if style is None:
        style = doc.styles.add_style(HEADING_STYLE_NAMES[level][0], WD_STYLE_TYPE.PARAGRAPH)
        base_style = find_style(doc, names=("Normal", "正文"), style_ids=("Normal",))
        if base_style is not None:
            style.base_style = base_style
    apply_style_format(style, f"heading_{level}")
    style.paragraph_format.line_spacing = line_spacing()
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    return style


def ensure_intro_style(doc: Document):
    style = find_style(doc, names=(INTRO_STYLE_NAME,))
    if style is None:
        style = doc.styles.add_style(INTRO_STYLE_NAME, WD_STYLE_TYPE.PARAGRAPH)
    base_style = find_style(doc, names=("Normal", "正文"), style_ids=("Normal",))
    if base_style is not None:
        style.base_style = base_style
    apply_style_format(style, "intro")
    style.paragraph_format.line_spacing = line_spacing()
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    ppr = style._element.get_or_add_pPr()
    for element_name in ("./w:outlineLvl", "./w:numPr"):
        for element in list(ppr.xpath(element_name)):
            ppr.remove(element)
    return style


def ensure_intro_subheading_style(doc: Document):
    style = find_style(doc, names=(INTRO_SUBHEADING_STYLE_NAME,))
    if style is None:
        style = doc.styles.add_style(INTRO_SUBHEADING_STYLE_NAME, WD_STYLE_TYPE.PARAGRAPH)
    base_style = find_style(doc, names=("Normal", "正文"), style_ids=("Normal",))
    if base_style is not None:
        style.base_style = base_style
    apply_style_format(style, "intro_subheading")
    style.paragraph_format.line_spacing = line_spacing()
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    ppr = style._element.get_or_add_pPr()
    for element_name in ("./w:outlineLvl", "./w:numPr"):
        for element in list(ppr.xpath(element_name)):
            ppr.remove(element)
    return style


def parse_numbered_heading(text: str):
    match = NUMBERED_HEADING_RE.match(text)
    if match is None:
        return None
    parts = [int(part) for part in match.group("number").split(".")]
    return parts, match.group("title")


def is_intro_heading_text(text: str) -> bool:
    if text == title_introduction():
        return True
    parsed = parse_numbered_heading(text)
    return parsed is not None and parsed[0] == [1] and parsed[1] == title_introduction()


def has_old_numbered_intro_structure(doc: Document, body_start: int) -> bool:
    first_text = clean(doc.paragraphs[body_start].text)
    parsed = parse_numbered_heading(first_text)
    if parsed is not None and parsed[0] == [1] and parsed[1] == title_introduction():
        return True
    if first_text != title_introduction():
        return False
    for paragraph in doc.paragraphs[body_start + 1 :]:
        parsed = parse_numbered_heading(clean(paragraph.text))
        if parsed is None:
            continue
        parts, _ = parsed
        return not (len(parts) == 1 and parts[0] == 1)
    return False


def format_numbered_heading(parts: list[int], title: str) -> str:
    return f"{'.'.join(str(part) for part in parts)} {title}"


def heading_level_from_text(text: str) -> int | None:
    if is_intro_heading_text(text):
        return 1
    parsed = parse_numbered_heading(text)
    if parsed is None:
        return None
    return len(parsed[0])


def first_body_index(doc: Document):
    front_end = find_paragraph(doc, title_tables())
    search_start = 0 if front_end is None else front_end + 1
    for i, paragraph in enumerate(doc.paragraphs):
        if i < search_start:
            continue
        text = clean(paragraph.text)
        if is_intro_heading_text(text):
            return i
        if is_heading_style(paragraph, levels=(1,)) and HEADING_TEXT_RE.match(text):
            return i
    for i, paragraph in enumerate(doc.paragraphs):
        if i < search_start:
            continue
        text = clean(paragraph.text)
        if is_intro_heading_text(text):
            return i
        if HEADING_TEXT_RE.match(text) and text not in front_titles():
            return i
    return None


def apply_heading_styles(doc: Document) -> int:
    body_start = first_body_index(doc)
    if body_start is None:
        return 0
    old_numbered_intro = has_old_numbered_intro_structure(doc, body_start)
    intro_style = ensure_intro_style(doc)
    intro_subheading_style = ensure_intro_subheading_style(doc)
    heading_styles = {
        1: get_heading_style(doc, 1),
        2: get_heading_style(doc, 2),
        3: get_heading_style(doc, 3),
    }
    changed = 0
    for paragraph in doc.paragraphs[body_start:]:
        text = clean(paragraph.text)
        if text in front_titles():
            continue
        level = heading_level_from_text(text)
        if level is None:
            continue
        parsed = parse_numbered_heading(text)
        if is_intro_heading_text(text):
            if text != title_introduction():
                replace_paragraph_text(paragraph, title_introduction())
            paragraph.style = intro_style
            clear_outline_level(paragraph)
        elif old_numbered_intro and parsed is not None and parsed[0][0] == 1:
            _, title = parsed
            replace_paragraph_text(paragraph, title)
            paragraph.style = intro_subheading_style
            clear_outline_level(paragraph)
        else:
            if old_numbered_intro and parsed is not None and parsed[0][0] > 1:
                parts, title = parsed
                shifted_parts = [parts[0] - 1, *parts[1:]]
                replace_paragraph_text(paragraph, format_numbered_heading(shifted_parts, title))
            paragraph.style = heading_styles[level]
            set_outline_level(paragraph, level)
        suppress_numbering(paragraph)
        changed += 1
    return changed


def make_toc_field(instr_text: str, placeholder: str):
    p = OxmlElement("w:p")
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:leader"), "dot")
    tab.set(qn("w:pos"), "8296")
    tabs.append(tab)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:line"), "360")
    spacing.set(qn("w:lineRule"), "auto")
    ppr.append(tabs)
    ppr.append(spacing)
    p.append(ppr)

    def add_child(child):
        r = OxmlElement("w:r")
        r.append(child)
        p.append(r)

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    add_child(begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instr_text
    add_child(instr)

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    add_child(separate)

    text = OxmlElement("w:t")
    text.text = placeholder
    add_child(text)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    add_child(end)
    return p


def ensure_front_title(doc: Document, text: str, page_break_before: bool = False):
    idx = find_paragraph(doc, text)
    if idx is not None:
        paragraph = doc.paragraphs[idx]
        normal_style = find_style(doc, names=("Normal", "正文"), style_ids=("a", "Normal"))
        if normal_style is not None:
            paragraph.style = normal_style
            set_paragraph_style_id(paragraph, getattr(normal_style, "style_id", "a") or "a")
        else:
            set_paragraph_style_id(paragraph, "a")
        clear_outline_level(paragraph)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.line_spacing = None if page_break_before else line_spacing()
        paragraph.paragraph_format.space_before = None
        paragraph.paragraph_format.space_after = None
        if page_break_before:
            ensure_page_break_at_start(paragraph)
        else:
            remove_page_breaks(paragraph)
        for run in paragraph.runs:
            apply_run_format(run, "front_title")
    return idx


def replace_between(doc: Document, start_text: str, end_text: str, replacement) -> bool:
    start = find_paragraph(doc, start_text)
    if start is None:
        return False
    end = find_paragraph(doc, end_text, start + 1)
    if end is None:
        return False
    for paragraph in list(doc.paragraphs[start + 1 : end]):
        remove_paragraph(paragraph)
    if replacement is not None:
        doc.paragraphs[start]._p.addnext(replacement)
    return True


def prepare_main_toc(doc: Document) -> bool:
    ensure_table_of_figures_style(doc)
    ensure_toc_styles(doc)
    ensure_front_title(doc, title_toc())
    ensure_front_title(doc, title_figures(), page_break_before=page_break_enabled("before_figures"))
    ensure_front_title(doc, title_tables(), page_break_before=page_break_enabled("before_tables"))
    return replace_between(
        doc,
        title_toc(),
        title_figures(),
        make_toc_field(r'TOC \o "1-3" \h \z \t "Report Intro Heading,1"', "目录将在 Word 中自动更新。"),
    )


def parse_toc_pages(doc: Document) -> dict[str, int]:
    pages = {}
    for paragraph in doc.paragraphs:
        if not style_name(paragraph).lower().startswith("toc "):
            continue
        text = clean(paragraph.text)
        intro_match = re.match(rf"^{re.escape(title_introduction())}\s+(\d+)$", text)
        if intro_match:
            pages[title_introduction()] = int(intro_match.group(1))
            continue
        match = TOC_LINE_RE.match(text)
        if not match:
            continue
        number, title, page = match.groups()
        pages[f"{number} {title}"] = int(page)
    return pages


def nearest_heading_page(heading_stack: list[str], pages: dict[str, int]):
    for heading in reversed(heading_stack):
        if heading in pages:
            return pages[heading]
    return None


def collect_caption_entries(doc: Document, pages: dict[str, int]):
    fig_style = ensure_style(doc, "图题注")
    tab_style = ensure_style(doc, "表题注")
    body_start = first_body_index(doc)
    if body_start is None:
        return [], []

    headings = []
    figures = []
    tables = []
    seen_figures = set()
    seen_tables = set()
    for paragraph in doc.paragraphs[body_start:]:
        text = clean(paragraph.text)
        if text == title_introduction() or is_heading_style(paragraph):
            headings.append(text)
            continue
        if FIG_RE.match(text):
            paragraph.style = fig_style
            if text not in seen_figures:
                figures.append((text, nearest_heading_page(headings, pages)))
                seen_figures.add(text)
        elif TAB_RE.match(text) and "从技术环节" not in text:
            paragraph.style = tab_style
            if text not in seen_tables:
                tables.append((text, nearest_heading_page(headings, pages)))
                seen_tables.add(text)
    return figures, tables


def write_static_list(doc: Document, start_text: str, end_text: str, entries) -> bool:
    list_style = ensure_table_of_figures_style(doc)
    start = find_paragraph(doc, start_text)
    if start is None:
        return False
    end = find_paragraph(doc, end_text, start + 1)
    if end is None:
        return False

    for paragraph in list(doc.paragraphs[start + 1 : end]):
        remove_paragraph(paragraph)

    anchor = doc.paragraphs[start]._p
    for caption, page in reversed(entries):
        p = doc.add_paragraph()
        p.style = list_style
        p.paragraph_format.left_indent = Twips(483)
        p.paragraph_format.first_line_indent = Twips(-485)
        p.paragraph_format.tab_stops.add_tab_stop(
            Twips(8296), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
        )
        p.paragraph_format.line_spacing = line_spacing()
        caption_run = p.add_run(caption)
        tab_run = p.add_run("\t")
        page_run = p.add_run("" if page is None else str(page))
        for run in (caption_run, tab_run, page_run):
            apply_run_format(run, "list_entry")
        anchor.addnext(p._p)
    return True


def first_body_heading_text(doc: Document) -> str | None:
    idx = first_body_index(doc)
    if idx is None:
        return None
    return clean(doc.paragraphs[idx].text)


def ensure_body_starts_on_new_page(doc: Document) -> bool:
    if not page_break_enabled("before_body"):
        return False
    idx = first_body_index(doc)
    if idx is None:
        return False
    ensure_page_break_at_start(doc.paragraphs[idx])
    return True


def prepare(path: Path, backup: bool) -> None:
    if backup:
        backup_path = path.with_name(f"{path.stem}.before-report-directories{path.suffix}")
        shutil.copy2(path, backup_path)
        print(f"backup={backup_path}")
    doc = Document(str(path))
    front_removed = remove_paragraphs_before(doc, title_toc())
    changed = apply_heading_styles(doc)
    toc_ok = prepare_main_toc(doc)
    doc.save(str(path))
    print(f"front_paragraphs_removed={front_removed}")
    print(f"heading_styles_applied={changed}")
    print(f"main_toc_prepared={toc_ok}")


def finalize(path: Path) -> None:
    doc = Document(str(path))
    pages = parse_toc_pages(doc)
    figures, tables = collect_caption_entries(doc, pages)
    body_heading = first_body_heading_text(doc)
    fig_ok = write_static_list(doc, title_figures(), title_tables(), figures)
    table_ok = False if body_heading is None else write_static_list(doc, title_tables(), body_heading, tables)
    body_page_break_ok = ensure_body_starts_on_new_page(doc)
    doc.save(str(path))
    print(f"toc_pages_found={len(pages)}")
    print(f"figures={len(figures)} figure_list_updated={fig_ok} missing_pages={sum(p is None for _, p in figures)}")
    print(f"tables={len(tables)} table_list_updated={table_ok} missing_pages={sum(p is None for _, p in tables)}")
    print(f"body_page_break_updated={body_page_break_ok}")


def main() -> None:
    parser = argparse.ArgumentParser(description="Prepare/update Word report directories.")
    parser.add_argument("path", type=Path)
    parser.add_argument("--phase", choices=["prepare", "finalize"], required=True)
    parser.add_argument("--no-backup", action="store_true")
    parser.add_argument("--config", type=Path)
    args = parser.parse_args()
    path = args.path.expanduser().resolve(strict=True)
    config_path = None if args.config is None else args.config.expanduser().resolve(strict=False)
    set_format_config(load_format_config(config_path))

    if args.phase == "prepare":
        prepare(path, backup=not args.no_backup)
    else:
        finalize(path)


if __name__ == "__main__":
    main()
